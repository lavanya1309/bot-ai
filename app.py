from flask import Flask, render_template, request, jsonify, send_file
import requests
import re
import markdown
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import HtmlFormatter
from datetime import datetime, timedelta
import json
import os
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import uuid

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'

# Groq AI API configuration
GROQ_API_KEY = "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"
GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
MODEL = "llama3-70b-8192"

# Directory to store chat history
CHAT_HISTORY_DIR = "chat_history"
DOCUMENT_DIR = "generated_documents"
os.makedirs(CHAT_HISTORY_DIR, exist_ok=True)
os.makedirs(DOCUMENT_DIR, exist_ok=True)

# Store the current conversation
current_conversation = []

def save_chat_history(history):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = os.path.join(CHAT_HISTORY_DIR, f"chat_{timestamp}.json")
    with open(filename, 'w') as f:
        json.dump(history, f)

def load_chat_history(filename):
    filepath = os.path.join(CHAT_HISTORY_DIR, filename)
    with open(filepath, 'r') as f:
        return json.load(f)

def get_last_day_chats():
    now = datetime.now()
    one_day_ago = now - timedelta(days=1)
    chat_files = [f for f in os.listdir(CHAT_HISTORY_DIR) if f.startswith("chat_") and f.endswith(".json")]
    last_day_chats = []
    for filename in sorted(chat_files, reverse=True):
        try:
            timestamp_str = filename[5:-5]
            chat_time = datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")
            if chat_time >= one_day_ago:
                history = load_chat_history(filename)
                if history:
                    summary = history[0]['query'][:50] + "..." if len(history[0]['query']) > 50 else history[0]['query']
                    last_day_chats.append({
                        'filename': filename,
                        'summary': summary,
                        'timestamp': chat_time.strftime("%H:%M")
                    })
        except ValueError:
            continue
    return last_day_chats

def format_code_blocks(text):
    def replace(match):
        language = match.group(1) or 'python'
        code = match.group(2)
        try:
            lexer = get_lexer_by_name(language)
        except:
            lexer = get_lexer_by_name('text')
        formatter = HtmlFormatter(style='monokai')
        highlighted_code = highlight(code, lexer, formatter)
        return f'<div class="code-block"><pre><code class="{language}">{highlighted_code}</code></pre></div>'
    
    pattern = re.compile(r'```(\w+)?\n(.*?)\n```', re.DOTALL)
    return re.sub(pattern, replace, text)

def format_tables(text):
    text = re.sub(r'\|\s*\|\s*\|', '| --- | --- |', text)

    def process_table(match):
        rows = [row.strip() for row in match.group(0).split('\n') if row.strip() and '---' not in row]
        if len(rows) < 2:
            return match.group(0)

        headers = [h.strip() for h in rows[0].split('|') if h.strip()]
        html = '<div class="table-container"><table class="comparison-table"><thead><tr>'
        html += ''.join(f'<th>{header}</th>' for header in headers) + '</tr></thead><tbody>'

        for row in rows[1:]:
            cells = [c.strip() for c in row.split('|') if c.strip()]
            if len(cells) == len(headers):
                html += '<tr>' + ''.join(f'<td>{cell}</td>' for cell in cells) + '</tr>'

        html += '</tbody></table></div>'
        return html

    text = re.sub(r'(\|.*\|.*\n)(\|.*\|\n)*', process_table, text)
    return text

def generate_document(content, doc_type):
    doc_id = str(uuid.uuid4())
    filename = f"document_{doc_id}.{doc_type}"
    filepath = os.path.join(DOCUMENT_DIR, filename)
    
    if doc_type == 'docx':
        doc = Document()
        doc.add_heading('Generated Document', 0)
        plain_text = re.sub(r'#+\s*', '', content)
        plain_text = re.sub(r'\*\*(.*?)\*\*', r'\1', plain_text)
        plain_text = re.sub(r'\*(.*?)\*', r'\1', plain_text)
        plain_text = re.sub(r'`(.*?)`', r'\1', plain_text)
        plain_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', plain_text)
        
        for paragraph in plain_text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        doc.save(filepath)
        
    elif doc_type == 'pdf':
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        
        plain_text = re.sub(r'#+\s*', '', content)
        plain_text = re.sub(r'\*\*(.*?)\*\*', r'\1', plain_text)
        plain_text = re.sub(r'\*(.*?)\*', r'\1', plain_text)
        plain_text = re.sub(r'`(.*?)`', r'\1', plain_text)
        plain_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', plain_text)
        
        story = []
        for paragraph in plain_text.split('\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph, styles["Normal"]))
        
        doc.build(story)
        
    elif doc_type == 'txt':
        plain_text = re.sub(r'#+\s*', '', content)
        plain_text = re.sub(r'\*\*(.*?)\*\*', r'\1', plain_text)
        plain_text = re.sub(r'\*(.*?)\*', r'\1', plain_text)
        plain_text = re.sub(r'`(.*?)`', r'\1', plain_text)
        plain_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', plain_text)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(plain_text)
    
    return filename

def generate_document_links(content):
    doc_types = ['docx', 'pdf', 'txt']
    links = []
    
    for doc_type in doc_types:
        filename = generate_document(content, doc_type)
        links.append({
            'type': doc_type,
            'url': f'/download_document/{filename}',
            'icon': 'file-word' if doc_type == 'docx' else 'file-pdf' if doc_type == 'pdf' else 'file-alt'
        })
    
    return links

@app.route('/')
def index():
    previous_chats = get_last_day_chats()
    formatted_history = []
    for item in current_conversation:
        formatted_query = markdown.markdown(item['query'])
        formatted_response = markdown.markdown(item['response'])
        formatted_response = format_code_blocks(formatted_response)
        formatted_response = format_tables(formatted_response)
        formatted_history.append({'query': formatted_query, 'response': formatted_response})
    return render_template('index.html', conversation=formatted_history, previous_chats=previous_chats)

@app.route('/ask', methods=['POST'])
def ask():
    query = request.form['query']
    if not query.strip():
        return jsonify({'error': 'Please enter a question'})

    # Check if user is asking for document generation
    document_requested = any(keyword in query.lower() for keyword in ['word document', 'pdf', 'text file', 'document', 'download', 'file'])

    system_prompt = """When a user asks for a document (Word, PDF, or TXT), provide the content in a clean format suitable for document generation. 
    Include a note that document download links have been provided."""
    
    messages = [
        {"role": "system", "content": system_prompt},
        *[
            {"role": "user" if i % 2 == 0 else "assistant", "content": msg['query'] if i % 2 == 0 else msg['response']}
            for i, msg in enumerate(current_conversation[-3:])
        ],
        {"role": "user", "content": query}
    ]

    try:
        response = requests.post(
            GROQ_URL,
            headers={
                "Authorization": f"Bearer {GROQ_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": MODEL,
                "messages": messages,
                "temperature": 0.7
            }
        )
        response.raise_for_status()
        reply = response.json()['choices'][0]['message']['content']
        
        # Generate document links if requested
        document_links = []
        if document_requested:
            document_links = generate_document_links(reply)
            reply += "\n\n**Document download links:**\n"
            for link in document_links:
                reply += f"- [{link['type'].upper()} Document](/download_document/{link['url'].split('/')[-1]})\n"
        
        current_conversation.append({'query': query, 'response': reply})

        formatted_reply = markdown.markdown(reply)
        formatted_reply = format_code_blocks(formatted_reply)
        formatted_reply = format_tables(formatted_reply)

        return jsonify({
            'response': formatted_reply, 
            'raw_response': reply,
            'document_links': document_links
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download_document/<filename>')
def download_document(filename):
    filepath = os.path.join(DOCUMENT_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    mimetypes = {
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pdf': 'application/pdf',
        'txt': 'text/plain'
    }
    
    file_ext = filename.split('.')[-1]
    return send_file(
        filepath,
        as_attachment=True,
        download_name=filename,
        mimetype=mimetypes.get(file_ext, 'application/octet-stream')
    )

@app.route('/new_chat', methods=['POST'])
def new_chat():
    global current_conversation
    if current_conversation:
        save_chat_history(current_conversation)
    current_conversation = []
    return jsonify({'success': True})

@app.route('/load_chat', methods=['POST'])
def load_chat():
    filename = request.form['filename']
    try:
        history = load_chat_history(filename)
        global current_conversation
        current_conversation = history
        
        formatted_history = []
        for item in history:
            formatted_query = markdown.markdown(item['query'])
            formatted_response = markdown.markdown(item['response'])
            formatted_response = format_code_blocks(formatted_response)
            formatted_response = format_tables(formatted_response)
            formatted_history.append({'query': formatted_query, 'response': formatted_response})
        
        return jsonify({'success': True, 'conversation': formatted_history})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/get_last_day_chats')
def get_last_day_chats_route():
    return jsonify(get_last_day_chats())

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
